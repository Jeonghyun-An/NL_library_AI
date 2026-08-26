"""
build_wikisource_manifest.py — 위키문헌(ko.wikisource.org)에서 저작권 만료 한국문학
작품을 받아 PDF로 만들고, upload_from_manifest.py가 바로 쓸 수 있는
manifest.jsonl을 생성한다.

절차: Special:Export로 원문(wikitext) 받기 → 마크업 제거 → docx 생성
      → LibreOffice로 PDF 변환 → manifest.jsonl 기록

사용:
    python build_wikisource_manifest.py --out-dir "C:/.../wikisource_books" \
        --soffice-script "<docx skill 경로>/scripts/office/soffice.py"
"""
import argparse
import io
import json
import re
import sys
import time

# Windows 콘솔(cp949)에서 한글/특수문자 print가 UnicodeEncodeError로 죽는 것 방지.
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
import urllib.error
import urllib.parse
import urllib.request
import subprocess
from pathlib import Path
from xml.etree import ElementTree as ET

import docx

# (book_id, 제목, 저자, 위키문헌 문서 제목) — 2026-08 확인 완료 목록
WORKS = [
    ("WS_001", "무정", "이광수", "무정"),
    ("WS_002", "진달래꽃", "김소월", "진달래꽃 (시집)"),
    ("WS_003", "님의 침묵", "한용운", "님의 침묵"),
    ("WS_004", "운수 좋은 날", "현진건", "운수 좋은 날"),
    ("WS_005", "메밀꽃 필 무렵", "이효석", "메밀꽃 필 무렵"),
    ("WS_006", "봄봄", "김유정", "봄봄"),
    ("WS_007", "태평천하", "채만식", "태평천하"),
    ("WS_008", "상록수", "심훈", "상록수"),
    ("WS_009", "인간문제", "강경애", "인간문제"),
    ("WS_010", "하늘과 바람과 별과 시", "윤동주", "하늘과 바람과 별과 시 (1948년)"),
    ("WS_011", "날개", "이상", "날개"),
    ("WS_012", "홍길동전", "허균", "홍길동전 (30장 경판본)"),
    ("WS_013", "흥부전", "저자미상", "흥부전 (경판 25장본)"),
    ("WS_014", "구운몽", "김만중", "구운몽"),
    ("WS_015", "벙어리 삼룡이", "나도향", "벙어리 삼룡이"),
    ("WS_016", "탈출기", "최서해", "탈출기"),
    ("WS_017", "광야", "이육사", "광야 (이육사)"),
    ("WS_018", "감자", "김동인", "감자"),
    ("WS_019", "배따라기", "김동인", "배따라기"),
    ("WS_020", "붉은 산", "김동인", "붉은 산"),
    ("WS_021", "정지용 시집", "정지용", "향수"),
    ("WS_022", "금수회의록", "안국선", "금수회의록"),
    ("WS_023", "혈의 누", "이인직", "혈의 누"),
    ("WS_024", "빈처", "현진건", "빈처"),
    ("WS_025", "고향", "현진건", "고향 (현진건)"),
    ("WS_026", "물레방아", "나도향", "물레방아"),
]

WIKISOURCE_EXPORT = "https://ko.wikisource.org/wiki/Special:Export/{}"
WIKISOURCE_API = "https://ko.wikisource.org/w/api.php"


def fetch_wikitext(page_title: str, *, max_retries: int = 5) -> str:
    """429(rate limit)는 지수 백오프로 재시도. 요청 사이 최소 간격도 둔다(REQUEST_DELAY)."""
    url = WIKISOURCE_EXPORT.format(urllib.parse.quote(page_title.replace(" ", "_")))
    req = urllib.request.Request(url, headers={"User-Agent": "nl-lib-ingest/1.0 (contact: ajh@landsoft.co.kr)"})
    delay = 2.0
    for attempt in range(max_retries):
        time.sleep(REQUEST_DELAY)
        try:
            with urllib.request.urlopen(req, timeout=30) as resp:
                xml_bytes = resp.read()
            break
        except urllib.error.HTTPError as e:
            if e.code == 429 and attempt < max_retries - 1:
                print(f"    429 — {delay:.0f}초 대기 후 재시도 ({page_title})")
                time.sleep(delay)
                delay *= 2
                continue
            raise
    else:
        raise RuntimeError(f"재시도 초과: {page_title}")

    root = ET.fromstring(xml_bytes)
    ns = {"mw": "http://www.mediawiki.org/xml/export-0.11/"}
    text_el = root.find(".//mw:page/mw:revision/mw:text", ns)
    if text_el is None or not text_el.text:
        raise RuntimeError(f"본문 없음: {page_title}")
    return text_el.text


REQUEST_DELAY = 1.2  # 위키문헌 요청 사이 최소 간격(초) — 429 예방


def list_subpages(page_title: str, *, max_retries: int = 5) -> list[str]:
    """장별로 하위 문서에 나뉜 장편 작품 대응 — <제목>/로 시작하는 하위 문서 목록."""
    params = {
        "action": "query", "list": "allpages",
        "apprefix": page_title + "/", "aplimit": "500", "format": "json",
    }
    url = WIKISOURCE_API + "?" + urllib.parse.urlencode(params)
    req = urllib.request.Request(url, headers={"User-Agent": "nl-lib-ingest/1.0 (contact: ajh@landsoft.co.kr)"})
    delay = 2.0
    for attempt in range(max_retries):
        time.sleep(REQUEST_DELAY)
        try:
            with urllib.request.urlopen(req, timeout=20) as resp:
                data = json.load(resp)
            return [p["title"] for p in data["query"]["allpages"]]
        except urllib.error.HTTPError as e:
            if e.code == 429 and attempt < max_retries - 1:
                time.sleep(delay)
                delay *= 2
                continue
            raise
    raise RuntimeError(f"재시도 초과(하위 문서 목록): {page_title}")


def _numeric_sort_key(subpage_title: str):
    """'제10장' 앞에 '제2장'이 오도록 하위 문서 이름의 선행 숫자로 정렬. 숫자가
    없으면(예: 시집의 개별 시 제목) 원문 순서를 못 얻으므로 이름순으로 대체한다."""
    tail = subpage_title.rsplit("/", 1)[-1]
    m = re.search(r"\d+", tail)
    return (0, int(m.group())) if m else (1, tail)


_PAGES_TAG = re.compile(
    r'<pages\s+index="([^"]+)"\s+from=(\d+)(?:\s+to=(\d+))?\s*/>'
)
_NOINCLUDE = re.compile(r"<noinclude>.*?</noinclude>", re.DOTALL)


def fetch_proofread_pages(index_name: str, start: int, end: int | None) -> str:
    """ProofreadPage 확장(<pages index=... from=X to=Y />)으로 스캔본을 트랜스클루전하는
    문서 대응 — 본문이 Special:Export에 안 잡히고 '페이지:<index>/<N>' 개별 문서에
    있다. to가 없으면 순번을 늘려가며 문서가 없어질 때까지 가져온다.
    """
    parts = []
    n = start
    while end is None or n <= end:
        title = f"페이지:{index_name}/{n}"
        try:
            wt = fetch_wikitext(title)
        except Exception as e:
            if end is None:
                break  # to 미지정 — 문서가 끊기는 지점을 끝으로 본다
            raise RuntimeError(f"스캔 페이지 실패({title}): {e}")
        parts.append(_NOINCLUDE.sub("", wt))
        n += 1
    if not parts:
        raise RuntimeError(f"스캔 페이지 없음: {index_name}")
    return "\n\n".join(parts)


def fetch_full_text(page_title: str) -> str:
    """본문이 하위 문서로 나뉜 장편은 각 장을 이어 붙이고, 아니면 본문서를 그대로 쓴다.

    하위 문서 중 하나라도 재시도 끝에 실패하면 전체를 실패 처리한다(일부만 붙여서
    "완성된 책"인 척 매니페스트에 넣는 게 더 나쁨 — 장이 통째로 빠진 걸 나중에
    알아채기 어렵다).
    """
    subpages = list_subpages(page_title)
    # '현대어 해석'/'번역' 류는 장 이어짐이 아니라 별개 버전 문서(예: 혈의 누/현대어
    # 해석) — 원전 본문과 이어붙이면 뒤섞이므로 하위 문서 취급에서 제외한다.
    subpages = [sp for sp in subpages if not re.search(r"현대어|번역", sp.rsplit("/", 1)[-1])]
    if subpages:
        subpages.sort(key=_numeric_sort_key)
        parts = []
        failed = []
        for sp in subpages:
            try:
                parts.append(fetch_wikitext(sp))
            except Exception as e:
                failed.append(sp)
                print(f"    하위 문서 실패({sp}): {e}")
        if failed:
            raise RuntimeError(
                f"하위 문서 {len(failed)}/{len(subpages)}건 실패 — 불완전한 본문이라 건너뜀: {failed}"
            )
        return "\n\n".join(parts)

    main_wt = fetch_wikitext(page_title)
    m = _PAGES_TAG.search(main_wt)
    if m:
        index_name, start, end = m.group(1), int(m.group(2)), m.group(3)
        return fetch_proofread_pages(index_name, start, int(end) if end else None)
    return main_wt


def wikitext_to_plain(wikitext: str) -> str:
    """위키 마크업 제거 — 완벽하지 않지만 본문 위주로 정리한다."""
    t = wikitext
    t = re.sub(r"<ref[^>]*/>", "", t)
    t = re.sub(r"<ref[^>]*>.*?</ref>", "", t, flags=re.DOTALL)
    t = re.sub(r"\{\{[^{}]*\}\}", "", t)               # 템플릿 {{...}}
    t = re.sub(r"\[\[(?:[^|\]]*\|)?([^\]]+)\]\]", r"\1", t)  # 링크 [[a|b]] -> b
    t = re.sub(r"'''('')?", "", t)                       # 굵게/기울임
    t = re.sub(r"^==+\s*(.*?)\s*==+$", r"\n\1\n", t, flags=re.MULTILINE)  # 절 제목
    t = re.sub(r"<[^>]+>", "", t)                        # 잔여 HTML 태그
    t = re.sub(r"^\s*[-*#:;].*$", "", t, flags=re.MULTILINE)  # 목차/표 기호 줄
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def build_docx(title: str, author: str, body: str, out_path: Path) -> None:
    d = docx.Document()
    d.add_heading(title, level=0)
    d.add_paragraph(f"저자: {author}")
    d.add_paragraph("")
    for para in body.split("\n\n"):
        para = para.strip()
        if para:
            d.add_paragraph(para)
    d.save(str(out_path))


def convert_to_pdf(soffice_bin: str, docx_path: Path, out_dir: Path) -> Path:
    subprocess.run(
        [soffice_bin, "--headless", "--convert-to", "pdf",
         "--outdir", str(out_dir), str(docx_path)],
        check=True, capture_output=True,
    )
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise RuntimeError(f"PDF 변환 실패: {docx_path}")
    return pdf_path


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out-dir", required=True, help="PDF·manifest를 저장할 디렉토리")
    ap.add_argument(
        "--soffice-bin",
        default=r"C:\Program Files\LibreOffice\program\soffice.exe",
        help="soffice(LibreOffice) 실행 파일 경로",
    )
    ap.add_argument("--bucket", default="nl-lib-bucket")
    ap.add_argument(
        "--only", default=None,
        help="쉼표로 구분한 book_id만 처리(예: WS_018,WS_019). 기존 manifest의 나머지 항목은 유지.",
    )
    args = ap.parse_args()

    out_dir = Path(args.out_dir)
    pdf_dir = out_dir / "pdf"
    docx_dir = out_dir / "_docx"
    pdf_dir.mkdir(parents=True, exist_ok=True)
    docx_dir.mkdir(parents=True, exist_ok=True)

    manifest_path = out_dir / "manifest.jsonl"
    only_ids = set(args.only.split(",")) if args.only else None
    works = [w for w in WORKS if only_ids is None or w[0] in only_ids]

    # 기존 manifest에서 이번에 다시 안 만드는 항목은 그대로 보존(머지).
    existing_rows = {}
    if manifest_path.exists():
        with open(manifest_path, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line:
                    r = json.loads(line)
                    existing_rows[r["book_id"]] = r
    processed_ids = {w[0] for w in works}
    rows = [r for bid, r in existing_rows.items() if bid not in processed_ids]
    new_ok = 0

    for book_id, title, author, page_title in works:
        print(f"[{book_id}] {title} ({author}) 처리 중...", flush=True)
        try:
            wikitext = fetch_full_text(page_title)
            plain = wikitext_to_plain(wikitext)
            if len(plain) < 200:
                print(f"  건너뜀 — 본문이 너무 짧음({len(plain)}자), 수동 확인 필요")
                continue

            docx_path = docx_dir / f"{book_id}.docx"
            build_docx(title, author, plain, docx_path)

            pdf_path = convert_to_pdf(args.soffice_bin, docx_path, pdf_dir)
            final_pdf = pdf_dir / f"{book_id}.pdf"
            if pdf_path != final_pdf:
                pdf_path.replace(final_pdf)

            size = final_pdf.stat().st_size
            rows.append({
                "book_id": book_id,
                "file": str(final_pdf),
                "object_key": f"originals/{book_id}/{book_id}.pdf",
                "size": size,
                "title": title,
            })
            new_ok += 1
            print(f"  완료 — {size:,} bytes")
        except Exception as e:
            print(f"  실패: {e}")

    with open(manifest_path, "w", encoding="utf-8") as f:
        for r in rows:
            f.write(json.dumps(r, ensure_ascii=False) + "\n")

    print(f"\n이번 처리 {new_ok}/{len(works)}건, 전체 manifest {len(rows)}건. manifest: {manifest_path}")


if __name__ == "__main__":
    main()
